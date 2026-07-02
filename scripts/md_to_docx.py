"""Convert the Week-8 submission markdown files into .docx for the LMS upload.

Lightweight handler — handles headings (#..####), bullets (- / *), numbered
lists (1.), code fences (```), inline `code`, **bold**, *italic*, and
pipe-delimited tables. Not a full CommonMark parser; it's tuned to OUR files.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "submission_docs"
OUT_DIR.mkdir(exist_ok=True)

FILES = [
    "WEEK8_SUBMISSION_README.md",
    "ARCHITECTURE.md",
    "PERFORMANCE_REPORT.md",
    "DEMO_VIDEO_SCRIPT.md",
    "FINAL_USER_CHECKLIST.md",
]


# ---------- inline formatting (bold / italic / code) ----------
_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def _add_runs(paragraph, text: str):
    """Split text into runs with **bold**, *italic*, `code` formatting."""
    parts = _INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def _flush_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        # Not a real table — emit as plain paragraphs
        for r in rows:
            doc.add_paragraph(" | ".join(r))
        return
    # Drop the separator row (---|---)
    body_rows = [rows[0]] + [r for r in rows[1:] if not all(re.fullmatch(r":?-+:?", c.strip()) for c in r)]
    cols = max(len(r) for r in body_rows)
    table = doc.add_table(rows=len(body_rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(body_rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, row[j] if j < len(row) else "")
            if i == 0:
                for run in p.runs:
                    run.bold = True


def convert(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)

    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()

        # Code fences
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_buf = []; in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Tables
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_buf.append(cells)
            continue
        else:
            flush_table()

        if not line.strip():
            doc.add_paragraph()
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", line.strip()):
            p = doc.add_paragraph()
            p.add_run("―" * 40)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1)); text = m.group(2)
            h = doc.add_heading(level=min(level, 4))
            _add_runs(h, text)
            continue

        # Bullets / numbered list
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(2))
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(2))
            continue

        # Block quote
        m = re.match(r"^>\s?(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(m.group(1))
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_runs(p, line)

    flush_table()
    doc.save(str(docx_path))


def main():
    for name in FILES:
        src = ROOT / name
        if not src.exists():
            print(f"skip (missing): {name}")
            continue
        dst = OUT_DIR / (src.stem + ".docx")
        convert(src, dst)
        print(f"wrote {dst.relative_to(ROOT)}  ({dst.stat().st_size//1024} KB)")


if __name__ == "__main__":
    main()
